"""Tests for the standalone document parser layer."""

from pathlib import Path

import pymupdf
import pytest
from docx import Document

from app.parsers import (
    DOCXParsingError,
    DocxParser,
    EmptyDocumentError,
    FileAccessError,
    MarkdownParser,
    PDFParser,
    PDFParsingError,
    ParserFactory,
    TextDecodingError,
    TextParser,
    UnsupportedFileTypeError,
)


def _create_pdf(path: Path) -> None:
    document = pymupdf.open()
    document.set_metadata({"title": "Research Notes", "author": "PaperForge"})
    document.new_page().insert_text((72, 72), "Introduction")
    document.new_page().insert_text((72, 72), "Findings")
    document.save(path)
    document.close()


def _create_docx(path: Path) -> None:
    document = Document()
    document.core_properties.title = "Research Notes"
    document.core_properties.author = "PaperForge"
    document.add_paragraph("First paragraph")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cell one"
    table.cell(0, 1).text = "Cell two"
    document.save(path)


def test_pdf_parser_extracts_pages_and_metadata(tmp_path: Path) -> None:
    file_path = tmp_path / "research.pdf"
    _create_pdf(file_path)

    parsed = PDFParser().parse(file_path)

    assert parsed.file_type == "pdf"
    assert parsed.page_count == 2
    assert "Introduction" in parsed.extracted_text
    assert "Findings" in parsed.extracted_text
    assert parsed.metadata["title"] == "Research Notes"
    assert parsed.metadata["author"] == "PaperForge"
    assert parsed.word_count == len(parsed.extracted_text.split())
    assert parsed.character_count == len(parsed.extracted_text)


def test_docx_parser_extracts_body_text_and_metadata(tmp_path: Path) -> None:
    file_path = tmp_path / "research.docx"
    _create_docx(file_path)

    parsed = DocxParser().parse(file_path)

    assert parsed.file_type == "docx"
    assert parsed.page_count is None
    assert "First paragraph" in parsed.extracted_text
    assert "Cell one\tCell two" in parsed.extracted_text
    assert parsed.metadata["title"] == "Research Notes"
    assert parsed.metadata["author"] == "PaperForge"


@pytest.mark.parametrize(
    ("parser", "filename", "content", "file_type"),
    [
        (MarkdownParser(), "notes.md", "# Heading\n\nBody text", "md"),
        (TextParser(), "notes.txt", "Plain text\nSecond line", "txt"),
    ],
)
def test_text_based_parsers_preserve_utf8_source(
    tmp_path: Path,
    parser: TextParser | MarkdownParser,
    filename: str,
    content: str,
    file_type: str,
) -> None:
    file_path = tmp_path / filename
    file_path.write_text(content, encoding="utf-8-sig", newline="")

    parsed = parser.parse(file_path)

    assert parsed.file_type == file_type
    assert parsed.extracted_text == content
    assert parsed.page_count is None
    assert parsed.metadata == {}
    assert parsed.word_count == len(content.split())
    assert parsed.character_count == len(content)


@pytest.mark.parametrize(
    ("file_path", "expected_type"),
    [
        (Path("document.PDF"), PDFParser),
        (Path("document.DOCX"), DocxParser),
        (Path("document.MD"), MarkdownParser),
        (Path("document.TXT"), TextParser),
    ],
)
def test_factory_selects_parser_case_insensitively(
    file_path: Path,
    expected_type: type[object],
) -> None:
    assert isinstance(ParserFactory.get_parser(file_path), expected_type)


def test_factory_rejects_unsupported_extension() -> None:
    with pytest.raises(UnsupportedFileTypeError, match="Unsupported document type"):
        ParserFactory.get_parser(Path("notes.csv"))


@pytest.mark.parametrize("content", ["", " \t\n"])
def test_empty_text_files_raise_meaningful_error(
    tmp_path: Path,
    content: str,
) -> None:
    file_path = tmp_path / "empty.txt"
    file_path.write_text(content, encoding="utf-8")

    with pytest.raises(EmptyDocumentError):
        TextParser().parse(file_path)


def test_invalid_utf8_text_raises_meaningful_error(tmp_path: Path) -> None:
    file_path = tmp_path / "invalid.txt"
    file_path.write_bytes(b"\xff\xfe\x00")

    with pytest.raises(TextDecodingError):
        TextParser().parse(file_path)


def test_missing_file_raises_meaningful_error(tmp_path: Path) -> None:
    with pytest.raises(FileAccessError, match="does not exist"):
        TextParser().parse(tmp_path / "missing.txt")


def test_corrupt_pdf_raises_meaningful_error(tmp_path: Path) -> None:
    file_path = tmp_path / "corrupt.pdf"
    file_path.write_bytes(b"this is not a PDF")

    with pytest.raises(PDFParsingError):
        PDFParser().parse(file_path)


def test_unreadable_docx_raises_meaningful_error(tmp_path: Path) -> None:
    file_path = tmp_path / "corrupt.docx"
    file_path.write_bytes(b"this is not a DOCX package")

    with pytest.raises(DOCXParsingError):
        DocxParser().parse(file_path)
