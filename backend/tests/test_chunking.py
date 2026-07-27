"""Tests for deterministic, paragraph-aware document chunking."""

from pathlib import Path

import pymupdf
import pytest
from docx import Document

from app.chunking import (
    ChunkingConfig,
    ChunkingInvariantError,
    ChunkingStrategy,
    DocumentChunker,
    EmptyDocumentError,
    InvalidChunkingConfigError,
    InvalidParsedDocumentError,
)
from app.chunking.strategy import ChunkSpan
from app.chunking.utils import count_words_and_content
from app.models.document_chunk import DocumentChunk
from app.models.parsed_document import ParsedDocument
from app.parsers import ParserFactory


def _document(
    text: str,
    *,
    filename: str = "research.txt",
    file_type: str = "txt",
    metadata: dict[str, str] | None = None,
) -> ParsedDocument:
    word_count, _ = count_words_and_content(text)
    return ParsedDocument(
        filename=filename,
        file_type=file_type,
        extracted_text=text,
        page_count=None,
        word_count=word_count,
        character_count=len(text),
        metadata=metadata or {},
    )


def _assert_source_spans(
    document: ParsedDocument,
    chunks: list[DocumentChunk],
) -> None:
    assert chunks[0].start_char == 0
    assert chunks[-1].end_char == len(document.extracted_text)

    previous = None
    reconstructed_parts: list[str] = []
    for index, chunk in enumerate(chunks):
        assert chunk.chunk_index == index
        assert chunk.text == document.extracted_text[chunk.start_char : chunk.end_char]
        assert chunk.character_count == len(chunk.text)
        assert chunk.word_count == len(chunk.text.split())
        if previous is not None:
            assert previous.start_char < chunk.start_char <= previous.end_char
            assert chunk.end_char > previous.end_char
            reconstructed_parts.append(
                chunk.text[previous.end_char - chunk.start_char :]
            )
        else:
            reconstructed_parts.append(chunk.text)
        previous = chunk

    assert "".join(reconstructed_parts) == document.extracted_text


def test_small_document_returns_structured_chunk() -> None:
    document = _document(
        "Research finding one.\n\nResearch finding two.",
        filename="notes.md",
        file_type="md",
        metadata={"author": "PaperForge"},
    )

    chunks = DocumentChunker().chunk(document)

    assert len(chunks) == 1
    chunk = chunks[0]
    assert chunk.document_filename == "notes.md"
    assert chunk.text == document.extracted_text
    assert chunk.metadata == {
        "source_file_type": "md",
        "source_metadata": {"author": "PaperForge"},
    }
    _assert_source_spans(document, chunks)


def test_chunks_have_exact_offsets_and_cover_large_document() -> None:
    paragraphs = [
        f"Paragraph {index} contains reproducible research evidence. " * 3
        for index in range(30)
    ]
    document = _document("\n\n".join(paragraphs))
    config = ChunkingConfig(max_chars=240, overlap_chars=40)

    chunks = DocumentChunker().chunk(document, config)

    assert len(chunks) > 2
    _assert_source_spans(document, chunks)


def test_overlap_is_exact_when_no_safe_boundary_exists() -> None:
    document = _document("x" * 250)
    chunks = DocumentChunker().chunk(
        document,
        ChunkingConfig(max_chars=100, overlap_chars=20),
    )

    assert [(chunk.start_char, chunk.end_char) for chunk in chunks] == [
        (0, 100),
        (80, 180),
        (160, 250),
    ]
    _assert_source_spans(document, chunks)


def test_minimum_useful_content_prevents_nested_overlap_chunks() -> None:
    intro = "Introduction evidence " * 24
    short_paragraph = "Brief transition."
    body = "Detailed research finding " * 80
    document = _document(f"{intro}\n\n{short_paragraph}\n\n{body}")
    config = ChunkingConfig(
        max_chars=500,
        overlap_chars=300,
        min_chunk_chars=160,
    )

    chunks = DocumentChunker().chunk(document, config)

    assert len(chunks) > 2
    for index, chunk in enumerate(chunks):
        if index < len(chunks) - 1:
            assert chunk.character_count <= config.max_chars
        else:
            assert chunk.character_count <= (
                config.max_chars + config.effective_min_chunk_chars
            )
        if index:
            assert chunk.end_char - chunks[index - 1].end_char >= (
                config.effective_min_chunk_chars
            )
    _assert_source_spans(document, chunks)


def test_short_final_tail_is_merged_within_final_size_limit() -> None:
    document = _document("x" * 200)
    config = ChunkingConfig(
        max_chars=100,
        overlap_chars=20,
        min_chunk_chars=30,
    )

    chunks = DocumentChunker().chunk(document, config)

    assert [(chunk.start_char, chunk.end_char) for chunk in chunks] == [
        (0, 100),
        (80, 200),
    ]
    assert chunks[0].character_count == config.max_chars
    assert config.max_chars < chunks[-1].character_count <= (
        config.max_chars + config.effective_min_chunk_chars
    )
    assert chunks[-1].end_char - chunks[0].end_char >= (
        config.effective_min_chunk_chars
    )
    _assert_source_spans(document, chunks)


def test_paragraphs_are_not_split_when_they_fit() -> None:
    paragraphs = [
        "First complete paragraph with supporting evidence.",
        "Second complete paragraph with supporting evidence.",
        "Third complete paragraph with supporting evidence.",
    ]
    document = _document("\n\n".join(paragraphs))

    chunks = DocumentChunker().chunk(
        document,
        ChunkingConfig(max_chars=70, overlap_chars=0),
    )

    for paragraph in paragraphs:
        assert any(paragraph in chunk.text for chunk in chunks)
    _assert_source_spans(document, chunks)


def test_markdown_heading_stays_with_following_paragraph() -> None:
    body = "The methods section explains the reproducible research process."
    document = _document(
        f"# Methods\n\n{body}\n\nNext independent section.",
        filename="paper.md",
        file_type="md",
    )

    chunks = DocumentChunker().chunk(
        document,
        ChunkingConfig(max_chars=80, overlap_chars=0),
    )

    heading_chunk = next(chunk for chunk in chunks if "# Methods" in chunk.text)
    assert body in heading_chunk.text


def test_oversized_heading_body_starts_with_its_heading() -> None:
    document = _document("# Results\n\n" + ("evidence " * 40))

    chunks = DocumentChunker().chunk(
        document,
        ChunkingConfig(max_chars=60, overlap_chars=0),
    )

    assert chunks[0].text.startswith("# Results")
    assert "evidence" in chunks[0].text
    _assert_source_spans(document, chunks)


def test_bullet_list_is_preserved_and_oversized_lists_split_by_item() -> None:
    items = [f"- Finding {index}: " + ("evidence " * 3) for index in range(6)]
    list_text = "\n".join(items)
    document = _document(list_text)

    chunks = DocumentChunker().chunk(
        document,
        ChunkingConfig(max_chars=80, overlap_chars=0),
    )

    for item in items:
        assert any(item in chunk.text for chunk in chunks)
    _assert_source_spans(document, chunks)


def test_extremely_long_paragraph_uses_deterministic_fallback_splits() -> None:
    document = _document("Sentence with research evidence. " * 30)
    config = ChunkingConfig(max_chars=75, overlap_chars=0)

    chunks = DocumentChunker().chunk(document, config)

    assert len(chunks) > 2
    assert all(chunk.character_count <= config.max_chars for chunk in chunks)
    _assert_source_spans(document, chunks)


def test_output_and_ids_are_deterministic() -> None:
    document = _document("Finding one.\n\nFinding two.\n\nFinding three.")
    config = ChunkingConfig(max_chars=20, overlap_chars=5)
    chunker = DocumentChunker()

    assert chunker.chunk(document, config) == chunker.chunk(document, config)


@pytest.mark.parametrize(
    "kwargs",
    [
        {"max_chars": 0, "overlap_chars": 0},
        {"max_chars": 10, "overlap_chars": -1},
        {"max_chars": 10, "overlap_chars": 10},
        {"max_chars": True, "overlap_chars": 0},
        {"max_chars": 10, "overlap_chars": 2, "min_chunk_chars": 0},
        {"max_chars": 10, "overlap_chars": 2, "min_chunk_chars": 9},
    ],
)
def test_invalid_chunking_configuration_is_rejected(kwargs: dict[str, int]) -> None:
    with pytest.raises(InvalidChunkingConfigError):
        ChunkingConfig(**kwargs)


def test_minimum_useful_content_configuration_is_deterministic() -> None:
    assert ChunkingConfig(
        max_chars=100,
        overlap_chars=20,
    ).effective_min_chunk_chars == 25
    assert ChunkingConfig(
        max_chars=100,
        overlap_chars=20,
        min_chunk_chars=30,
    ).effective_min_chunk_chars == 30


def test_empty_and_whitespace_documents_are_rejected() -> None:
    empty_document = ParsedDocument.model_construct(
        filename="empty.txt",
        file_type="txt",
        extracted_text="",
        page_count=None,
        word_count=0,
        character_count=0,
        metadata={},
    )

    with pytest.raises(EmptyDocumentError):
        DocumentChunker().chunk(empty_document)
    with pytest.raises(EmptyDocumentError):
        DocumentChunker().chunk(_document(" \t\n"))


def test_corrupted_parsed_document_is_rejected() -> None:
    document = _document("Reliable evidence")
    mismatched_count = document.model_copy(update={"word_count": 1})
    invalid_metadata = document.model_copy(
        update={"metadata": {"nested": {"not": "allowed"}}}
    )

    with pytest.raises(InvalidParsedDocumentError, match="word_count"):
        DocumentChunker().chunk(mismatched_count)
    with pytest.raises(InvalidParsedDocumentError, match="metadata"):
        DocumentChunker().chunk(invalid_metadata)


class _BrokenStrategy(ChunkingStrategy):
    def plan_spans(
        self,
        text: str,
        config: ChunkingConfig,
    ):
        del config
        yield ChunkSpan(0, min(4, len(text)))
        yield ChunkSpan(0, len(text))


def test_invalid_strategy_output_is_rejected() -> None:
    with pytest.raises(ChunkingInvariantError, match="non-progressing"):
        DocumentChunker(_BrokenStrategy()).chunk(_document("Evidence remains."))


def test_parser_factory_integration_for_markdown_pdf_and_docx(tmp_path: Path) -> None:
    markdown_path = tmp_path / "research.md"
    markdown_path.write_text("# Title\n\nMarkdown evidence.", encoding="utf-8")

    pdf_path = tmp_path / "research.pdf"
    pdf = pymupdf.open()
    pdf.new_page().insert_text((72, 72), "PDF evidence")
    pdf.save(pdf_path)
    pdf.close()

    docx_path = tmp_path / "research.docx"
    docx = Document()
    docx.add_paragraph("DOCX evidence")
    docx.save(docx_path)

    chunker = DocumentChunker()
    for path in (markdown_path, pdf_path, docx_path):
        document = ParserFactory.parse(path)
        chunks = chunker.chunk(document, ChunkingConfig(max_chars=100, overlap_chars=0))
        assert chunks[0].document_filename == path.name
        assert chunks[0].metadata["source_file_type"] == document.file_type
        _assert_source_spans(document, chunks)


def test_hundreds_of_paragraphs_have_stable_linear_span_invariants() -> None:
    paragraph = "Research evidence supports a reproducible conclusion. " * 10
    document = _document("\n\n".join(paragraph for _ in range(300)))

    chunks = DocumentChunker().chunk(
        document,
        ChunkingConfig(max_chars=500, overlap_chars=50),
    )

    assert len(chunks) > 100
    _assert_source_spans(document, chunks)
